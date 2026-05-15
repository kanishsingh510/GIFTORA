from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(249,115,22)
def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(30,30,30)
def para(t, bold=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(30,30,30)
def bullet(t, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if prefix:
        r1 = p.add_run(prefix+": "); r1.bold = True; r1.font.size = Pt(11)
        p.add_run(t).font.size = Pt(11)
    else:
        p.add_run(t).font.size = Pt(11)
def code_block(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)

# ═══════════════════ COVER ═══════════════════
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("GIFTORA STUDIO"); tr.bold = True; tr.font.size = Pt(32)
tr.font.color.rgb = RGBColor(249,115,22)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Technical Setup & Installation Guide"); sr.bold = True; sr.font.size = Pt(18)
doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
subr = sub.add_run("Milestone 1: Environment & Project Initialization"); subr.italic = True; subr.font.size = Pt(13)

doc.add_page_break()

h1("1. Prerequisites")
para("Ensure you have the following installed on your local development machine:")
bullet("Node.js (v18.x or later)", prefix="Runtime")
bullet("npm (v9.x or later)", prefix="Package Manager")
bullet("MongoDB (Local) or MongoDB Atlas (Cloud account)", prefix="Database")
bullet("Git", prefix="Version Control")

h1("2. Project Structure")
para("Giftora is a MERN stack application with a decoupled architecture:")
bullet("React 18 + Vite (Client-side application)", prefix="/client")
bullet("Node.js + Express (Server-side API)", prefix="/server")

h1("3. Installation Steps")
h2("3.1 Clone and Install")
code_block("""git clone https://github.com/siddhantkumar101/GIFTORA.git
cd GIFTORA

# Install all dependencies (both client and server)
npm run install:all""")

h2("3.2 Backend Configuration")
para("Navigate to the server directory and create a .env file:")
code_block("""cd server
cp .env.example .env""")
para("Update the .env file with your credentials:")
code_block("""PORT=5000
CLIENT_ORIGIN=http://localhost:5173
MONGODB_URI=your_mongodb_connection_string
JWT_SECRET=your_secure_secret_key
RAZORPAY_KEY_ID=your_razorpay_key_id
RAZORPAY_KEY_SECRET=your_razorpay_key_secret""")

h2("3.3 Frontend Configuration")
para("The frontend is configured to proxy requests to localhost:5000 in development. No manual .env is required for local setup unless deploying to production.")

h1("4. Running the Application")
para("From the root directory, run the development script:")
code_block("npm run dev")
para("This will concurrently start:")
bullet("Frontend: http://localhost:5173")
bullet("Backend: http://localhost:5000")

h1("5. Database Seeding")
para("To populate your database with the initial 15 products for Milestone 1:")
code_block("""cd server
npm run seed""")

h1("6. Common Commands")
bullet("Compiles frontend for deployment", prefix="npm run build (in /client)")
bullet("Runs the API server separately", prefix="npm start (in /server)")
bullet("Standard dev mode", prefix="npm run dev")

doc.save(r"C:\Users\siddh\Desktop\Giftora_Setup_Guide.docx")
print("Setup Guide saved!")
