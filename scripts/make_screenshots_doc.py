from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(249,115,22)

def para(t):
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(12)

# ═══════════════════ COVER ═══════════════════
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("GIFTORA STUDIO"); tr.bold = True; tr.font.size = Pt(32)
tr.font.color.rgb = RGBColor(249,115,22)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Production Screenshots & Features"); sr.bold = True; sr.font.size = Pt(18)
doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
subr = sub.add_run("Live Website Showcase: https://giftora-six.vercel.app"); subr.italic = True

doc.add_page_break()

# We'll use the artifacts paths from the prompt history
images = [
    ("Homepage Hero & Header", r"C:\Users\siddh\.gemini\antigravity\brain\a8f80ef0-8601-410a-bf72-8e11988fd2e8\homepage_full_1778830078842.png"),
    ("Product Categories & Catalog", r"C:\Users\siddh\.gemini\antigravity\brain\a8f80ef0-8601-410a-bf72-8e11988fd2e8\homepage_categories_1778830205775.png"),
    ("Mobile Optimized Layout", r"C:\Users\siddh\.gemini\antigravity\brain\a8f80ef0-8601-410a-bf72-8e11988fd2e8\homepage_hero_1778830191452.png"),
]

for title, path in images:
    h1(title)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6))
    else:
        para(f"[Image not found at {path} - Please replace with a real screenshot of the {title}]")
    doc.add_page_break()

doc.save(r"C:\Users\siddh\Desktop\Giftora_Screenshots.docx")
print("Screenshots document saved!")
